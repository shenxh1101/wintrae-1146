import os
from datetime import datetime
from typing import List, Optional
from sqlalchemy.orm import Session
from sqlalchemy import and_
from models import Invoice, VerificationStatus, ReviewStatus
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.units import inch


class ExportService:
    @staticmethod
    def export_to_excel(db: Session, invoices: List[Invoice], output_path: str) -> str:
        wb = Workbook()
        ws = wb.active
        ws.title = "发票查验记录"
        
        headers = [
            "发票代码", "发票号码", "发票类型", "开票日期",
            "购买方名称", "购买方税号",
            "销售方名称", "销售方税号",
            "金额", "税额", "价税合计",
            "验证状态", "复核状态", "费用类别",
            "是否黑名单", "黑名单原因",
            "异常原因", "提交人ID", "提交时间"
        ]
        
        ws.append(headers)
        
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)
        
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        for invoice in invoices:
            row = [
                invoice.invoice_code or "",
                invoice.invoice_number or "",
                invoice.invoice_type.value if invoice.invoice_type else "",
                invoice.invoice_date.strftime("%Y-%m-%d") if invoice.invoice_date else "",
                invoice.buyer_name or "",
                invoice.buyer_tax_number or "",
                invoice.seller_name or "",
                invoice.seller_tax_number or "",
                invoice.amount_without_tax or 0,
                invoice.tax_amount or 0,
                invoice.total_amount or 0,
                invoice.verification_status.value if invoice.verification_status else "",
                invoice.review_status.value if invoice.review_status else "",
                invoice.expense_category or "",
                "是" if invoice.is_blacklisted else "否",
                invoice.blacklisted_reason or "",
                invoice.exception_reason or "",
                invoice.submitter_id,
                invoice.created_at.strftime("%Y-%m-%d %H:%M:%S") if invoice.created_at else ""
            ]
            ws.append(row)
        
        for col_num in range(1, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col_num)].width = 18
        
        wb.save(output_path)
        return output_path

    @staticmethod
    def export_to_word(db: Session, invoices: List[Invoice], output_path: str) -> str:
        doc = Document()
        
        title = doc.add_heading("发票查验报告", 0)
        title.alignment = 1
        
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"发票总数: {len(invoices)}")
        
        total_amount = sum(inv.total_amount or 0 for inv in invoices)
        total_tax = sum(inv.tax_amount or 0 for inv in invoices)
        verified_count = sum(1 for inv in invoices if inv.verification_status == VerificationStatus.VERIFIED)
        suspicious_count = sum(1 for inv in invoices if inv.verification_status == VerificationStatus.SUSPICIOUS)
        
        doc.add_heading("汇总统计", level=1)
        doc.add_paragraph(f"价税合计总额: ¥{total_amount:,.2f}")
        doc.add_paragraph(f"税额总额: ¥{total_tax:,.2f}")
        doc.add_paragraph(f"验证通过: {verified_count} 张")
        doc.add_paragraph(f"可疑发票: {suspicious_count} 张")
        
        doc.add_heading("明细列表", level=1)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        headers = ['发票号码', '销售方', '金额', '验证状态', '复核状态']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        for invoice in invoices[:100]:
            row_cells = table.add_row().cells
            row_cells[0].text = invoice.invoice_number or ""
            row_cells[1].text = invoice.seller_name or ""
            row_cells[2].text = f"¥{invoice.total_amount:,.2f}" if invoice.total_amount else "¥0.00"
            row_cells[3].text = invoice.verification_status.value if invoice.verification_status else ""
            row_cells[4].text = invoice.review_status.value if invoice.review_status else ""
        
        doc.save(output_path)
        return output_path

    @staticmethod
    def export_to_pdf(db: Session, invoices: List[Invoice], output_path: str) -> str:
        doc = SimpleDocTemplate(output_path, pagesize=landscape(A4))
        
        elements = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1
        )
        
        elements.append(Paragraph("发票查验报告", title_style))
        elements.append(Spacer(1, 0.2 * inch))
        
        summary_text = f"""
        生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>
        发票总数: {len(invoices)}<br/>
        价税合计: ¥{sum(inv.total_amount or 0 for inv in invoices):,.2f}<br/>
        税额合计: ¥{sum(inv.tax_amount or 0 for inv in invoices):,.2f}
        """
        elements.append(Paragraph(summary_text, styles['Normal']))
        elements.append(Spacer(1, 0.3 * inch))
        
        data = [['序号', '发票号码', '销售方', '金额', '税额', '验证状态', '复核状态']]
        
        for idx, invoice in enumerate(invoices[:50], 1):
            data.append([
                str(idx),
                invoice.invoice_number or "",
                invoice.seller_name or "",
                f"¥{invoice.total_amount:,.2f}" if invoice.total_amount else "¥0.00",
                f"¥{invoice.tax_amount:,.2f}" if invoice.tax_amount else "¥0.00",
                invoice.verification_status.value if invoice.verification_status else "",
                invoice.review_status.value if invoice.review_status else ""
            ])
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(table)
        doc.build(elements)
        
        return output_path
